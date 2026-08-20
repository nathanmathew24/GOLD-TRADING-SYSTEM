"""
generate_docs.py — Word Document Generator

Creates the system_explanation.docx documentation file using python-docx.
Run this file once to generate the document:
    python docs/generate_docs.py
"""

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level=1):
    """Add a styled heading to the document."""
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)  # Dark navy
    return heading


def add_body(doc, text):
    """Add a normal paragraph."""
    para = doc.add_paragraph(text)
    para.style.font.size = Pt(11)
    return para


def add_bullet(doc, text):
    """Add a bullet point."""
    return doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc, text):
    """Add a numbered list item."""
    return doc.add_paragraph(text, style="List Number")


def build_document():
    doc = Document()

    # ── Document Title ────────────────────────────────────────────────────────
    title = doc.add_heading("Regime-Aware Multi-Agent Gold Trading System", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Technical Documentation & Judge Reference Guide")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.italic = True
    doc.add_paragraph()

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 1: PROJECT OVERVIEW
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "1. Project Overview", 1)
    add_body(doc,
        "This system is an automated, AI-powered trading assistant for the Gold (XAU/USD) market. "
        "It works by first understanding what the market is currently doing — is it trending strongly, "
        "moving erratically, or just drifting sideways? Once it knows the 'regime', it activates the right "
        "specialist agent to analyze the market and produce a trading recommendation."
    )
    add_body(doc,
        "The system uses four cooperating AI agents that each have a specific job. They pass information to "
        "each other in a pipeline, and the final output is a clean, human-readable report printed in the "
        "terminal showing exactly what trade to consider, where to enter, where to put a stop loss, and "
        "how large the position should be."
    )
    add_body(doc,
        "The entire system runs from a single command: python main.py"
    )

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 2: WHY GOLD
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "2. Why Gold?", 1)
    add_body(doc,
        "We chose Gold (XAU/USD) for several strong reasons:"
    )
    add_bullet(doc, "Safe Haven Asset: Gold is globally recognized as a store of value. It reacts predictably to macroeconomic events like inflation, interest rate changes, and geopolitical crises.")
    add_bullet(doc, "High Liquidity: The gold market trades around the clock and has extremely deep liquidity, meaning price data is accurate and manipulation is difficult.")
    add_bullet(doc, "Technical Clarity: Gold responds well to technical analysis. It forms clear trends, recognizable patterns, and respects support and resistance levels.")
    add_bullet(doc, "Data Availability: Reliable, free historical data is available via Yahoo Finance (GC=F ticker), making it ideal for automated systems.")
    add_bullet(doc, "Regime Variety: Gold naturally transitions between trending, ranging, and volatile states — making it a perfect test case for a regime-aware system.")

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 3: WHAT IS A MARKET REGIME
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "3. What Is a Market Regime?", 1)
    add_body(doc,
        "A market regime describes the overall 'personality' or 'mode' that a market is currently operating in. "
        "Just like weather has different states (sunny, stormy, cloudy), markets have different states too. "
        "Using the same strategy in different regimes would be like wearing a raincoat on a sunny day — technically "
        "possible, but not optimal."
    )
    add_body(doc, "Our system identifies three regimes:")

    add_heading(doc, "TRENDING", 2)
    add_body(doc,
        "The market is moving strongly and consistently in one direction — either up or down. "
        "Detected when ADX > 25 and price is clearly above or below the 50-day EMA. "
        "Best strategy: follow the trend with momentum indicators."
    )

    add_heading(doc, "VOLATILE", 2)
    add_body(doc,
        "The market is making unusually large price swings — more than 1.5x its normal daily movement range. "
        "This could be due to news events, data releases, or sudden shocks. "
        "High risk environment — our system defaults to NEUTRAL in this regime."
    )

    add_heading(doc, "RANGING", 2)
    add_body(doc,
        "The market has no clear trend and is oscillating between support and resistance levels. "
        "ADX is low and volatility is normal. "
        "Best strategy: mean-reversion (buy low, sell high within the range). Our system defaults to NEUTRAL here."
    )

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 4: AGENT BREAKDOWN
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "4. Agent Breakdown", 1)

    agents = [
        (
            "Agent 1: Orchestrator (orchestrator.py)",
            "The master coordinator. It runs each agent in sequence, passes data between them, and handles errors.",
            "None — it self-starts the pipeline.",
            "Coordinates everything; produces the final assembled result set.",
        ),
        (
            "Agent 2: Regime Classifier (regime_agent.py)",
            "Fetches 90 days of Gold price data and classifies the current market into TRENDING, VOLATILE, or RANGING using ADX, ATR, and EMA50.",
            "None (fetches its own data from Yahoo Finance).",
            "Regime label (e.g., TRENDING), confidence score, and key metric values.",
        ),
        (
            "Agent 3: Momentum Agent (momentum_agent.py)",
            "Activates only in TRENDING regimes. Analyzes EMA crossover, MACD, and RSI to determine if the trend is LONG (upward) or SHORT (downward).",
            "Regime result from Agent 2.",
            "Trade signal (LONG/SHORT/NEUTRAL), entry price zone, and indicator readings.",
        ),
        (
            "Agent 4: Risk Agent (risk_agent.py)",
            "Takes the trade signal and calculates safe trade parameters. Uses the 2% portfolio risk rule and ATR-based stop placement.",
            "Momentum result from Agent 3.",
            "Stop loss price, take profit price, position size in ounces, maximum dollar risk.",
        ),
        (
            "Agent 5: Report Agent (report_agent.py)",
            "Makes a single OpenAI API call to generate a plain-English trade thesis. Renders a formatted, color-coded report in the terminal using the Rich library.",
            "All three previous agent results.",
            "Final human-readable report printed to terminal.",
        ),
    ]

    for name, role, inputs, outputs in agents:
        add_heading(doc, name, 2)
        add_body(doc, f"Role: {role}")
        add_body(doc, f"Input: {inputs}")
        add_body(doc, f"Output: {outputs}")
        doc.add_paragraph()

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 5: THE FLOW
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "5. The Flow — Step by Step", 1)
    add_body(doc, "Here is exactly what happens when you run python main.py:")

    steps = [
        "main.py starts and calls the Orchestrator.",
        "Orchestrator calls the Regime Agent, which downloads 90 days of Gold price data from Yahoo Finance.",
        "Regime Agent computes ADX, ATR, and EMA50. It classifies the market as TRENDING, VOLATILE, or RANGING and returns this result.",
        "Orchestrator reads the regime. If TRENDING, it activates the Momentum Agent. If VOLATILE or RANGING, it creates a default NEUTRAL signal.",
        "Momentum Agent (if activated) computes EMA9/21 crossover, MACD, and RSI. It outputs a LONG, SHORT, or NEUTRAL signal with entry zone.",
        "Risk Agent takes the signal and current price. It calculates: stop loss (1.5x ATR away), take profit (2:1 reward), and position size (2% risk rule).",
        "Report Agent takes all results and calls the OpenAI GPT-4o-mini API to write a 3-4 sentence trade thesis.",
        "Report Agent uses the Rich library to render a color-coded, formatted report in the terminal with all trade parameters.",
    ]
    for step in steps:
        add_numbered(doc, step)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 6: KEY CONCEPTS EXPLAINED
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "6. Key Concepts Explained", 1)

    concepts = [
        (
            "ADX — Average Directional Index",
            "ADX measures the STRENGTH of a trend, not its direction. The scale is 0 to 100. "
            "Below 20: weak or no trend. 20–25: developing trend. Above 25: strong trend. Above 40: very strong trend. "
            "ADX doesn't tell you if the trend is up or down — just how strong it is. We pair it with EMA50 to get direction."
        ),
        (
            "ATR — Average True Range",
            "ATR measures VOLATILITY — how much the price typically moves each day. "
            "A higher ATR means larger daily swings. We use ATR in two ways: "
            "(1) to detect volatile regimes (when ATR is 1.5x its 20-day average), and "
            "(2) to set stop losses that are calibrated to actual market movement rather than arbitrary dollar amounts."
        ),
        (
            "EMA Crossover Strategy",
            "EMA stands for Exponential Moving Average — it smooths price data while giving more weight to recent prices. "
            "We use two EMAs: a fast one (9-period) and a slow one (21-period). "
            "When the fast EMA crosses above the slow EMA, it signals upward momentum (LONG signal). "
            "When the fast EMA crosses below the slow EMA, it signals downward momentum (SHORT signal)."
        ),
        (
            "MACD — Moving Average Convergence Divergence",
            "MACD measures the relationship between two moving averages (12-period and 26-period EMA). "
            "The MACD line crossing above zero = bullish momentum. Below zero = bearish. "
            "The histogram shows how fast momentum is building or fading. A rising histogram confirms a developing trend."
        ),
        (
            "RSI — Relative Strength Index",
            "RSI (14-period) measures how fast and by how much price has been moving recently. Scale: 0–100. "
            "Above 70 = overbought (price may be due for a pullback — risky to enter LONG). "
            "Below 30 = oversold (price may bounce — risky to enter SHORT). "
            "We use RSI as a filter — if RSI contradicts our signal, we downgrade to NEUTRAL."
        ),
        (
            "Risk-Reward Ratio",
            "The risk-reward ratio compares how much we stand to lose vs. how much we stand to gain. "
            "A 2:1 ratio means for every $1 we risk, we aim to make $2. "
            "This ensures that even if only 40% of trades win, the system remains profitable overall. "
            "We calculate: risk distance (from entry to stop loss) and target distance (2x the risk distance)."
        ),
        (
            "Position Sizing — The 2% Rule",
            "The 2% rule is a professional risk management principle: never risk more than 2% of your total portfolio on a single trade. "
            "Formula: Position Size = (Portfolio × 2%) / Stop Distance. "
            "Example: $100,000 portfolio × 2% = $2,000 max risk. If stop is $20 away, we can buy 100 oz. "
            "This keeps any single loss small enough that the account can recover."
        ),
    ]

    for title_text, explanation in concepts:
        add_heading(doc, title_text, 2)
        add_body(doc, explanation)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 7: QUESTIONS JUDGES MIGHT ASK
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "7. Questions We Might Be Asked", 1)

    qa_pairs = [
        (
            "Why did you choose a multi-agent architecture instead of a single script?",
            "Multi-agent systems are modular — each agent has one job and does it well. This makes the system "
            "easier to debug, extend, and explain. If we want to add a mean-reversion strategy for ranging markets, "
            "we just add a new agent without touching the others. Single scripts become 'spaghetti code' fast."
        ),
        (
            "What happens when the market is VOLATILE or RANGING?",
            "Currently, the system outputs a NEUTRAL signal and doesn't recommend a trade. This is intentional — "
            "it is better to miss an opportunity than to take a bad trade. Future versions would route to "
            "specialized agents: a mean-reversion agent for ranging markets, and a volatility breakout agent for volatile ones."
        ),
        (
            "How accurate is the regime classification?",
            "The regime classification is based on well-established technical principles (ADX, ATR, EMA) used by "
            "professional traders. It is not 100% accurate — no indicator is. But the combination of three different "
            "lenses (trend strength, volatility level, price relative to EMA) gives a robust multi-dimensional view."
        ),
        (
            "Why gpt-4o-mini and not a larger model?",
            "GPT-4o-mini is used only for generating the plain-English trade thesis — a relatively simple writing task. "
            "All the actual analysis and signal generation is done by deterministic Python code, not the LLM. "
            "Using a smaller model keeps costs low and latency fast. The LLM is a narrator, not an analyst."
        ),
        (
            "Is this system actually profitable? Would you trade with it?",
            "This is a demonstration and educational system. In a production environment, you would need: "
            "backtesting across historical data, walk-forward analysis, transaction cost modeling, slippage simulation, "
            "and paper trading before live deployment. The framework is sound, but no trading system should go live "
            "without extensive validation."
        ),
        (
            "What does the 2% rule protect against?",
            "The 2% rule protects against catastrophic drawdowns. If you risk 2% per trade and have 10 losing trades "
            "in a row (very unlikely), you've only lost ~18% of your account (compounding). At 10% risk per trade, "
            "10 losses would wipe out ~65% of the account. Professional risk management keeps losses survivable."
        ),
        (
            "How does ATR-based stop placement work?",
            "Instead of placing a stop at an arbitrary number like '$50', we base it on the market's actual volatility. "
            "If gold moves $20/day on average, a $10 stop would get triggered by normal noise. "
            "By using 1.5x ATR, we place the stop wide enough to survive normal market breathing room, "
            "but tight enough to limit our loss if the trade goes wrong."
        ),
        (
            "Why use EMA instead of SMA (Simple Moving Average)?",
            "EMA gives more weight to recent prices, making it more responsive to current market conditions. "
            "A simple moving average treats a price from 50 days ago the same as yesterday's price, which can lag "
            "behind real market movements. For trading signals, being responsive to recent data is more valuable."
        ),
        (
            "Could this system be used for other assets?",
            "Yes. The architecture is asset-agnostic. You would only need to change the ticker in config.py (e.g., "
            "'BTC-USD' for Bitcoin, 'SPY' for S&P 500 ETF). The indicators, regime logic, and risk rules all apply "
            "universally to any liquid, tradeable asset with OHLCV data available on Yahoo Finance."
        ),
        (
            "What are the biggest limitations of this system?",
            "Key limitations: (1) No backtesting — we don't know the historical win rate. (2) Only one active strategy "
            "(momentum) — VOLATILE and RANGING regimes produce no signal. (3) No real-time data — uses daily bars, "
            "not tick data. (4) The LLM thesis is generated, not verified — it should be read as a summary, not gospel."
        ),
    ]

    for i, (question, answer) in enumerate(qa_pairs, 1):
        add_heading(doc, f"Q{i}: {question}", 2)
        add_body(doc, answer)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 8: GLOSSARY
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "8. Glossary", 1)

    glossary = {
        "ADX (Average Directional Index)": "A technical indicator measuring the strength of a price trend on a scale of 0–100. Values above 25 indicate a trending market.",
        "ATR (Average True Range)": "A volatility indicator measuring the average daily price range over a set period (typically 14 days).",
        "EMA (Exponential Moving Average)": "A type of moving average that gives more weight to recent prices, making it more responsive to new information.",
        "MACD (Moving Average Convergence Divergence)": "A trend-following momentum indicator that shows the relationship between two exponential moving averages.",
        "RSI (Relative Strength Index)": "A momentum oscillator (0–100) that measures the speed and change of price movements. Above 70 = overbought, below 30 = oversold.",
        "Market Regime": "The overall 'state' or 'mode' a market is operating in — typically classified as trending, volatile, or ranging.",
        "LONG": "A trade that profits when the price goes UP. You buy first, then sell later at a higher price.",
        "SHORT": "A trade that profits when the price goes DOWN. You sell first (borrowing the asset), then buy back later at a lower price.",
        "Stop Loss": "A pre-set price level at which a losing trade is automatically closed to limit further losses.",
        "Take Profit": "A pre-set price level at which a winning trade is automatically closed to lock in gains.",
        "Risk-Reward Ratio": "The ratio of potential profit to potential loss on a trade. A 2:1 ratio means you aim to make $2 for every $1 risked.",
        "Position Sizing": "Calculating how many units of an asset to trade so that a losing trade only costs a predetermined percentage of total portfolio value.",
        "2% Rule": "A risk management principle stating you should never risk more than 2% of your total trading capital on a single trade.",
        "GC=F": "Yahoo Finance ticker symbol for Gold Futures (COMEX). Tracks the spot price of gold.",
        "XAU/USD": "The forex pair notation for Gold priced in US Dollars. XAU is the international currency code for Gold.",
        "OHLCV": "Open, High, Low, Close, Volume — the five standard data points recorded for each period in a price chart.",
        "EMA Crossover": "A trading signal generated when a faster EMA crosses above or below a slower EMA, indicating a potential trend change.",
        "Agent": "In AI systems, an autonomous software component that perceives its environment, makes decisions, and takes actions to achieve a specific goal.",
        "Multi-Agent System": "A system composed of multiple AI agents that each handle distinct tasks and communicate their results to each other.",
        "Pipeline": "A series of data processing steps where the output of one step becomes the input of the next.",
        "OpenAI API": "The API provided by OpenAI that allows developers to access GPT language models programmatically.",
        "GPT-4o-mini": "A fast, cost-efficient version of OpenAI's GPT-4o language model, suitable for text generation tasks.",
    }

    for term, definition in glossary.items():
        para = doc.add_paragraph()
        run_term = para.add_run(f"{term}: ")
        run_term.bold = True
        run_term.font.size = Pt(11)
        run_def = para.add_run(definition)
        run_def.font.size = Pt(11)

    # ── Save the document ─────────────────────────────────────────────────────
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "system_explanation.docx")
    doc.save(output_path)
    print(f"[DocGen] Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    build_document()
